"""
generate_changelog.py
──────────────────────
Generates the project changelog as a formatted .docx file.
Run this script whenever you want to update the changelog document.

Usage:
    python generate_changelog.py

Output:
    PG_Recommendation_Changelog.docx  (in the project root)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─────────────────────────────────────────────
# STYLING HELPERS
# ─────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set table cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),  "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc, text: str, level: int, color_hex: str = "1F3864"):
    """Add a styled heading with custom colour."""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    return heading


def add_bullet(doc, text: str, bold_part: str = None):
    """Add a bullet point, optionally bolding the first part."""
    para = doc.add_paragraph(style="List Bullet")
    if bold_part and text.startswith(bold_part):
        run = para.add_run(bold_part)
        run.bold = True
        para.add_run(text[len(bold_part):])
    else:
        para.add_run(text)
    return para


def add_table(doc, headers: list, rows: list, header_color: str = "1F3864"):
    """Add a formatted table with coloured header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = header
        set_cell_bg(cell, header_color)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = "F2F2F2" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, cell_val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(cell_val)
            set_cell_bg(cell, bg)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    return table


# ─────────────────────────────────────────────
# DOCUMENT CONTENT
# ─────────────────────────────────────────────

def build_changelog() -> Document:
    doc = Document()

    # ── Page margins ─────────────────────────
    section = doc.sections[0]
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

    # ── Title Block ───────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("🏠 PG Accommodation Recommendation System")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Project Change Log & Developer Notes")
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    sub_run.italic = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"Document Version: 1.0   |   "
        f"Last Updated: {datetime.date.today().strftime('%d %B %Y')}   |   "
        f"Author: Parth Jha"
    ).font.size = Pt(9)

    doc.add_paragraph()   # spacer

    # ═══════════════════════════════════════
    # SECTION 1 — PURPOSE OF THIS DOCUMENT
    # ═══════════════════════════════════════
    add_heading(doc, "1. Purpose of This Document", level=1)
    doc.add_paragraph(
        "This document provides a complete, plain-language record of every change "
        "made to the PG Accommodation Recommendation System project. "
        "It is written so that anyone — including non-technical readers — can understand "
        "what was changed, why it was changed, and how the system works after the change. "
        "A new entry will be added each time the project is updated."
    )

    # ═══════════════════════════════════════
    # SECTION 2 — PROJECT OVERVIEW
    # ═══════════════════════════════════════
    add_heading(doc, "2. Project Overview", level=1)
    doc.add_paragraph(
        "The PG Accommodation Recommendation System helps students and working professionals "
        "find the best-matching PG (Paying Guest) accommodation based on their preferences "
        "such as rent budget, WiFi, AC, meals, gender, and location. "
        "It uses a lightweight scoring approach (no heavy machine learning) to rank available "
        "PGs from a dataset and show the top matches with a rating and explanation."
    )

    add_heading(doc, "Technology Stack", level=2)
    add_bullet(doc, "Python 3.10+")
    add_bullet(doc, "Pandas & NumPy — data processing")
    add_bullet(doc, "scikit-learn MinMaxScaler — feature normalization")
    add_bullet(doc, "rapidfuzz — fuzzy location matching (new)")
    add_bullet(doc, "python-docx — this changelog document (new)")

    # ═══════════════════════════════════════
    # SECTION 3 — VERSION LOG
    # ═══════════════════════════════════════
    add_heading(doc, "3. Change Log", level=1)

    # ────────────────────────────────────────
    # Version 1.0 — Initial Monolith
    # ────────────────────────────────────────
    add_heading(doc, "Version 1.0  —  Initial Working System", level=2)
    doc.add_paragraph(
        f"Date: 13 April 2026   |   File: main.py (single file, 194 lines)"
    ).runs[0].italic = True

    doc.add_paragraph(
        "This was the original version of the system. Everything — data loading, "
        "cleaning, encoding, user input, scoring, and output — was written in a single "
        "Python file called main.py. The system worked but had several important bugs "
        "and limitations described below."
    )

    add_heading(doc, "What Was Working in Version 1.0", level=3)
    add_bullet(doc, "Loaded PG data from CSV file")
    add_bullet(doc, "Cleaned and encoded Yes/No columns, Gender, Food Type, Curfew")
    add_bullet(doc, "Applied MinMaxScaler to Rent, Meals, and Floors")
    add_bullet(doc, "Accepted user input from command line")
    add_bullet(doc, "Hard-filtered PGs by availability and budget")
    add_bullet(doc, "Scored PGs using weighted absolute difference from user preferences")
    add_bullet(doc, "Converted scores to ratings and displayed a ranked table")

    add_heading(doc, "Known Bugs in Version 1.0 (Now Fixed)", level=3)

    bug_rows = [
        ["#", "Bug Description", "Where in Code", "Impact"],
        ["1", "Rent comparison mismatch — budget was normalized manually but "
              "PG rent was scaled by MinMaxScaler using a different formula, "
              "so the comparison was unreliable.",
         "Lines 102–122",
         "HIGH — wrong PGs could be filtered in/out"],
        ["2", "Gender filter too strict — used exact equality, so Co-ed PGs "
              "were never shown to Boys or Girls users.",
         "Line 122",
         "HIGH — missed valid PGs"],
        ["3", "Meals per day scale mismatch — user entered meals/3 (raw division) "
              "while PG values were already scaled by MinMaxScaler.",
         "Lines 113, 144",
         "MEDIUM — scoring inaccurate"],
        ["4", "Location match all-or-nothing — only gave a -0.2 bonus for "
              "exact string match; any typo or partial name got zero credit.",
         "Lines 147–148",
         "MEDIUM — poor location ranking"],
        ["5", "No empty-result guard — if no PGs matched the filters, the "
              "program crashed with a confusing pandas error instead of a "
              "friendly message.",
         "Line 172",
         "MEDIUM — poor user experience"],
        ["6", "Rating could be 0.0 — the normalization formula produced "
              "a rating of 0.0 for the worst match instead of 1.0.",
         "Lines 160–161",
         "LOW — misleading output"],
        ["7", "9 out of 14 encoded features were never used in scoring — "
              "Security, CCTV, Power Backup, Housekeeping, Parking, Lift, "
              "Drinking Water, Curfew Time, Food Type were all encoded but ignored.",
         "Lines 133–140",
         "MEDIUM — scoring incomplete"],
        ["8", "No output explanation — users had no idea why a PG was recommended.",
         "Lines 177–194",
         "LOW — poor transparency"],
        ["9", "Hardcoded weights — users could not express personal priorities.",
         "Lines 133–140",
         "LOW — not personalized"],
        ["10","All logic in one file — no separation of concerns; impossible "
              "to test individual components or extend without breaking things.",
         "Entire file",
         "HIGH — bad architecture"],
    ]

    # Build bug table manually (skip the header row from data)
    headers = bug_rows[0]
    rows    = bug_rows[1:]
    add_table(doc, headers, rows)

    doc.add_paragraph()   # spacer

    # ────────────────────────────────────────
    # Version 2.0 — Refactored & Bug-Fixed
    # ────────────────────────────────────────
    add_heading(doc, "Version 2.0  —  Full Refactor & Bug Fixes", level=2)
    doc.add_paragraph(
        f"Date: 14 April 2026   |   Files changed: main.py (rewritten), "
        f"config.py (new), src/data_loader.py (new), src/preprocessor.py (new), "
        f"src/recommender.py (new), src/explainer.py (new)"
    ).runs[0].italic = True

    doc.add_paragraph(
        "This version completely restructures the project from one file into separate "
        "modules, fixes all 10 known bugs, and adds new intelligence features. "
        "The result is a system that is easier to read, test, maintain, and extend."
    )

    # ── 2a. Architectural Changes ──
    add_heading(doc, "2a. Architectural Changes — What Was Reorganised", level=3)
    doc.add_paragraph(
        "The single 194-line main.py has been split into 6 focused files:"
    )

    arch_rows = [
        ["File", "What It Does", "Why It Exists Separately"],
        ["config.py",
         "Stores all settings: file paths, weights, thresholds, gender rules",
         "One place to change any setting without touching logic code"],
        ["src/data_loader.py",
         "Reads the CSV, checks all columns exist, removes duplicates",
         "Keeps file I/O isolated; easy to swap CSV for a database later"],
        ["src/preprocessor.py",
         "Encodes Yes/No columns, scales numbers, engineers new features",
         "Preprocessing is reusable and testable on its own"],
        ["src/recommender.py",
         "Applies hard filter, scores PGs, assigns ratings and ranks",
         "The brain of the system; can be called by CLI or a web API"],
        ["src/explainer.py",
         "Generates 'Why this PG?' text and assigns badges to results",
         "Explanation logic is separate so it can be styled for any UI"],
        ["main.py (rewritten)",
         "Orchestrates everything: asks user, calls modules, prints results",
         "Acts as the glue; contains no business logic itself"],
    ]
    add_table(doc, arch_rows[0], arch_rows[1:])

    doc.add_paragraph()

    add_heading(doc, "New Folder Structure", level=3)
    folder_para = doc.add_paragraph()
    folder_para.add_run(
        "PG_Recommendation/\n"
        "├── data/\n"
        "│   ├── raw/   ← CSV file lives here\n"
        "│   └── processed/   ← future: cached pre-processed data\n"
        "├── src/\n"
        "│   ├── __init__.py\n"
        "│   ├── data_loader.py\n"
        "│   ├── preprocessor.py\n"
        "│   ├── recommender.py\n"
        "│   └── explainer.py\n"
        "├── config.py\n"
        "├── main.py\n"
        "├── requirements.txt\n"
        "└── PG_Recommendation_Changelog.docx"
    ).font.name = "Courier New"

    doc.add_paragraph()

    # ── 2b. Bug Fixes ──
    add_heading(doc, "2b. Bug Fixes — What Was Fixed and How", level=3)

    fix_rows = [
        ["Bug #", "Fix Applied", "Simple Explanation"],
        ["1 — Rent comparison",
         "Hard filter now uses Original_Rent (raw ₹) column which is saved "
         "before MinMaxScaler runs. Budget (₹) is compared directly to Original_Rent (₹).",
         "We now compare apples to apples — both values in rupees."],
        ["2 — Gender filter",
         "A GENDER_COMPAT dictionary maps each user preference to a list of "
         "acceptable PG genders. Boys users see Boys + Co-ed. Girls see Girls + Co-ed.",
         "Co-ed PGs are now correctly shown to boys and girls users."],
        ["3 — Meals scale",
         "A normalize_user_meals() helper applies the same formula "
         "MinMaxScaler uses: (meals - 2) / 1, giving 0.0 for 2 meals and 1.0 for 3.",
         "User meals and PG meals are now on the same 0-to-1 scale."],
        ["4 — Location matching",
         "rapidfuzz.fuzz.partial_ratio() returns a 0–100 similarity score "
         "between the user's location text and each PG's location. "
         "Divided by 100, this is used as a continuous location bonus.",
         "Partial names and typos now still get credit. 'Hennur' matches 'Hennur Road'."],
        ["5 — Empty result crash",
         "recommend() now checks if df_filtered is empty and raises a "
         "friendly ValueError with helpful tips instead of crashing.",
         "Users now see a helpful message instead of a Python stack trace."],
        ["6 — Rating floor at 0",
         "New formula: rating = (1 - normalized) * 4 + 1,  giving [1.0, 5.0] range.",
         "The worst match now shows 1.0 star, not 0.0 stars."],
        ["7 — Missing features in scoring",
         "Security, Housekeeping, and Power_Backup added to DEFAULT_WEIGHTS in config.py.",
         "More features now contribute to the match score."],
        ["8 — No explanation",
         "explainer.py adds a 'Why_Matched' column with icons and reasons "
         "and a 'Badge' column (Budget Pick, Premium, Best Value, Available Now).",
         "Each recommendation now shows why it was chosen."],
        ["9 — Hardcoded weights",
         "All weights live in config.py DEFAULT_WEIGHTS dict. "
         "Any caller can pass a custom weights dict to override them.",
         "Weights can be tuned in one place without touching algorithm code."],
        ["10 — Monolith architecture",
         "Entire codebase split into 6 files with clear responsibilities.",
         "Each module can be read, tested, and modified independently."],
    ]
    add_table(doc, fix_rows[0], fix_rows[1:])

    doc.add_paragraph()

    # ── 2c. New Features ──
    add_heading(doc, "2c. New Features Added", level=3)

    feat_rows = [
        ["Feature", "What It Does", "File"],
        ["amenity_score",
         "Counts how many of the 10 key amenities a PG has, as a fraction (0 to 1). "
         "A score of 0.8 means 8 out of 10 amenities are available.",
         "src/preprocessor.py"],
        ["price_to_value",
         "Calculates rent divided by amenity_score. A lower number means "
         "you get more amenities for less money — a better deal.",
         "src/preprocessor.py"],
        ["available_soon",
         "Flags PGs that become available within the next 7 days as '1' (yes).",
         "src/preprocessor.py"],
        ["curfew_score",
         "Converts curfew time into a 0-to-1 number: 0 = strictest (9 PM), "
         "1 = no curfew. Used for display and future scoring.",
         "src/preprocessor.py"],
        ["explain_match()",
         "Generates a human-readable sentence explaining why a PG was recommended. "
         "Example: ✅ WiFi | 💰 Well within budget | 🔒 Security + CCTV",
         "src/explainer.py"],
        ["assign_badge()",
         "Assigns one badge to each PG: 🟢 Budget Pick, ⭐ Premium, "
         "💡 Best Value, or 🔥 Available Now.",
         "src/explainer.py"],
        ["Input validation",
         "All user inputs now have while-loop guards that re-ask on invalid entry "
         "instead of crashing with a Python error.",
         "main.py"],
        ["Location autocomplete hint",
         "Before asking for location, the system shows a list of available "
         "locations from the dataset so users know what to type.",
         "main.py"],
        ["Rich card display",
         "Results are now shown as formatted cards with icons, amenity lists, "
         "ratings, badges, and explanation — not just a raw data table.",
         "main.py"],
        ["Fuzzy location matching",
         "Uses rapidfuzz partial_ratio to give continuous location similarity "
         "scores instead of on/off exact match bonus.",
         "src/recommender.py"],
    ]
    add_table(doc, feat_rows[0], feat_rows[1:])

    doc.add_paragraph()

    # ── 2d. Algorithm Improvement ──
    add_heading(doc, "2d. Algorithm Improvement — Scoring Formula", level=3)

    doc.add_paragraph(
        "The original scoring formula used a simple sum of absolute differences "
        "(Manhattan distance). The new formula uses Weighted Euclidean Distance, "
        "which is mathematically more reliable because it squares the difference "
        "before summing — this means large mismatches are penalised more heavily "
        "than small mismatches."
    )

    doc.add_paragraph("Old formula (Version 1.0):")
    old_f = doc.add_paragraph()
    old_f.add_run("score = Σ ( weight_i × |pg_i − user_i| )").font.name = "Courier New"

    doc.add_paragraph("New formula (Version 2.0):")
    new_f = doc.add_paragraph()
    new_f.add_run(
        "distance_score = √ Σ ( weight_i × (pg_i − user_i)² )\n"
        "location_bonus = LOCATION_WEIGHT × fuzzy_ratio(pg_location, user_location)\n"
        "final_score    = distance_score − location_bonus\n\n"
        "match_rating   = ( (1 − normalized_score) × 4 ) + 1    [always in 1.0 – 5.0]"
    ).font.name = "Courier New"

    doc.add_paragraph(
        "Lower final_score = better PG match. "
        "The location bonus reduces the score when the PG location is similar to "
        "what the user asked for (so a good location improves the ranking)."
    )

    # ═══════════════════════════════════════
    # SECTION 4 — UPCOMING ROADMAP
    # ═══════════════════════════════════════
    add_heading(doc, "4. Upcoming Improvements (Planned)", level=1)

    road_rows = [
        ["Priority", "Feature", "Description", "Version Target"],
        ["P1", "Flask REST API",
         "Wrap the recommender in a web API: POST /recommend, GET /locations",
         "v2.1"],
        ["P1", "HTML/JS Card UI",
         "A browser-based form with result cards (no React, just HTML + CSS + JS)",
         "v2.1"],
        ["P2", "Dynamic weights from user",
         "Let user drag-rank their priorities; weights adjust automatically",
         "v2.2"],
        ["P2", "NLP text input",
         "Parse natural speech like 'cheap PG near Hennur with WiFi' into filters",
         "v2.2"],
        ["P3", "PostgreSQL migration",
         "Move data from CSV to a real database; add indexes on rent, location, gender",
         "v3.0"],
        ["P3", "Redis caching",
         "Cache repeated (location + budget) queries for 5 minutes",
         "v3.0"],
        ["Future", "KNN recommendation",
         "Once 500+ user sessions are collected, use K-Nearest Neighbours",
         "v4.0"],
        ["Future", "Clustering PGs",
         "Auto-group PGs into Budget / Mid / Premium tiers using KMeans",
         "v4.0"],
    ]
    add_table(doc, road_rows[0], road_rows[1:])

    doc.add_paragraph()

    # ═══════════════════════════════════════
    # SECTION 5 — HOW TO RUN
    # ═══════════════════════════════════════
    add_heading(doc, "5. How to Run the Project", level=1)

    doc.add_paragraph("Step 1: Install dependencies")
    s1 = doc.add_paragraph()
    s1.add_run("pip install -r requirements.txt").font.name = "Courier New"

    doc.add_paragraph("Step 2: Run the recommendation system")
    s2 = doc.add_paragraph()
    s2.add_run("python main.py").font.name = "Courier New"

    doc.add_paragraph("Step 3: Regenerate this changelog document")
    s3 = doc.add_paragraph()
    s3.add_run("python generate_changelog.py").font.name = "Courier New"

    doc.add_paragraph()

    # ═══════════════════════════════════════
    # FOOTER NOTE
    # ═══════════════════════════════════════
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        f"Generated automatically on {datetime.datetime.now().strftime('%d %B %Y, %I:%M %p')}  "
        f"·  PG Recommendation System  ·  Parth Jha"
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    footer_run.italic = True

    return doc


# ─────────────────────────────────────────────
# ENTRY POINT
# ─────────────────────────────────────────────

if __name__ == "__main__":
    output_path = "PG_Recommendation_Changelog.docx"
    print("📄  Generating changelog document...")
    doc = build_changelog()
    doc.save(output_path)
    print(f"✅  Saved: {output_path}")
